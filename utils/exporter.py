"""
AI小说写作助手 - 导出模块
支持导出为 txt, docx, pdf 格式
"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


class NovelExporter:
    """小说导出器"""

    def export_txt(self, title, content, export_dir):
        """导出为TXT文件"""
        filename = f"{title}.txt"
        filepath = os.path.join(export_dir, filename)

        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)

        return filepath

    def export_docx(self, title, full_content, characters, plot, export_dir):
        """导出为Word文档"""
        filename = f"{title}.docx"
        filepath = os.path.join(export_dir, filename)

        doc = Document()

        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = 'SimSun'
        font.size = Pt(12)

        # 标题
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(f"《{title}》")
        title_run.font.size = Pt(22)
        title_run.bold = True
        title_run.font.name = 'SimHei'

        doc.add_paragraph()

        # 角色介绍
        if characters:
            heading = doc.add_paragraph()
            heading_run = heading.add_run('【角色介绍】')
            heading_run.font.size = Pt(14)
            heading_run.bold = True
            heading_run.font.name = 'SimHei'

            for char in characters:
                if isinstance(char, dict):
                    name = char.get('name', '')
                    desc_parts = []
                    if char.get('role'):
                        desc_parts.append(f"身份：{char['role']}")
                    if char.get('personality'):
                        desc_parts.append(f"性格：{char['personality']}")
                    if char.get('appearance'):
                        desc_parts.append(f"外貌：{char['appearance']}")
                    if char.get('description'):
                        desc_parts.append(f"描述：{char['description']}")

                    char_text = f"  {name}"
                    if desc_parts:
                        char_text += "：" + "，".join(desc_parts)

                    p = doc.add_paragraph(char_text)
                    p.paragraph_format.left_indent = Cm(0.5)
                else:
                    p = doc.add_paragraph(f"  {char}")
                    p.paragraph_format.left_indent = Cm(0.5)

            doc.add_paragraph()

        # 情节大纲
        if plot:
            heading = doc.add_paragraph()
            heading_run = heading.add_run('【情节大纲】')
            heading_run.font.size = Pt(14)
            heading_run.bold = True
            heading_run.font.name = 'SimHei'

            p = doc.add_paragraph(plot)
            p.paragraph_format.left_indent = Cm(0.5)

            doc.add_paragraph()

        # 正文标题
        body_heading = doc.add_paragraph()
        body_heading_run = body_heading.add_run('【正文】')
        body_heading_run.font.size = Pt(14)
        body_heading_run.bold = True
        body_heading_run.font.name = 'SimHei'

        doc.add_paragraph()

        # 正文内容（按段落分割）
        # 分离正文部分
        if '【正文】' in full_content:
            body_start = full_content.index('【正文】') + len('【正文】')
            body_text = full_content[body_start:].strip()
            # 跳过分隔线
            lines = body_text.split('\n')
            content_started = False
            for line in lines:
                if not content_started:
                    if line.strip() == '=' * 40:
                        content_started = True
                    continue
                p = doc.add_paragraph(line.strip())
                p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进两字符
        else:
            for line in full_content.split('\n'):
                if line.strip():
                    p = doc.add_paragraph(line.strip())
                    p.paragraph_format.first_line_indent = Cm(0.74)

        doc.save(filepath)
        return filepath

    def export_pdf(self, title, full_content, characters, plot, export_dir):
        """导出为PDF文件"""
        filename = f"{title}.pdf"
        filepath = os.path.join(export_dir, filename)

        from fpdf import FPDF

        class ChinesePDF(FPDF):
            def __init__(self):
                super().__init__()
                # 尝试使用系统自带中文字体
                self.font_path = self._find_chinese_font()

            def _find_chinese_font(self):
                """查找系统中文字体"""
                # Windows常见中文字体路径
                possible_paths = [
                    'C:/Windows/Fonts/simsun.ttc',
                    'C:/Windows/Fonts/simhei.ttf',
                    'C:/Windows/Fonts/msyh.ttc',
                    'C:/Windows/Fonts/msyhbd.ttc',
                    'C:/Windows/Fonts/simkai.ttf',
                    # Linux路径
                    '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc',
                    '/usr/share/fonts/truetype/wqy/wqy-microhei.ttc',
                    '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
                    # macOS路径
                    '/System/Library/Fonts/PingFang.ttc',
                    '/System/Library/Fonts/STHeiti Light.ttc',
                ]
                for path in possible_paths:
                    if os.path.exists(path):
                        return path
                return None

        pdf = ChinesePDF()
        pdf.set_auto_page_break(auto=True, margin=15)

        # 注册中文字体
        if pdf.font_path:
            pdf.add_font('CJK', '', pdf.font_path, uni=True)
            pdf.add_font('CJK', 'B', pdf.font_path, uni=True)
            font_name = 'CJK'
        else:
            font_name = 'Helvetica'

        pdf.add_page()

        # 标题
        pdf.set_font(font_name, 'B', 20)
        pdf.cell(0, 15, f'《{title}》', align='C', new_x="LMARGIN", new_y="NEXT")
        pdf.ln(10)

        # 角色介绍
        if characters:
            pdf.set_font(font_name, 'B', 13)
            pdf.cell(0, 10, '【角色介绍】', new_x="LMARGIN", new_y="NEXT")

            pdf.set_font(font_name, '', 11)
            for char in characters:
                if isinstance(char, dict):
                    name = char.get('name', '')
                    desc_parts = []
                    if char.get('role'):
                        desc_parts.append(f"身份：{char['role']}")
                    if char.get('personality'):
                        desc_parts.append(f"性格：{char['personality']}")
                    if char.get('appearance'):
                        desc_parts.append(f"外貌：{char['appearance']}")
                    if char.get('description'):
                        desc_parts.append(f"描述：{char['description']}")

                    char_text = f"  {name}"
                    if desc_parts:
                        char_text += "：" + "，".join(desc_parts)
                    pdf.cell(0, 7, char_text, new_x="LMARGIN", new_y="NEXT")
                else:
                    pdf.cell(0, 7, f"  {char}", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(5)

        # 情节大纲
        if plot:
            pdf.set_font(font_name, 'B', 13)
            pdf.cell(0, 10, '【情节大纲】', new_x="LMARGIN", new_y="NEXT")
            pdf.set_font(font_name, '', 11)
            pdf.multi_cell(0, 7, plot)
            pdf.ln(5)

        # 正文
        pdf.set_font(font_name, 'B', 13)
        pdf.cell(0, 10, '【正文】', new_x="LMARGIN", new_y="NEXT")
        pdf.ln(5)

        pdf.set_font(font_name, '', 11)

        # 提取正文部分
        body_text = full_content
        if '【正文】' in full_content:
            body_start = full_content.index('【正文】') + len('【正文】')
            body_text = full_content[body_start:].strip()
            lines = body_text.split('\n')
            content_started = False
            for line in lines:
                if not content_started:
                    if line.strip() == '=' * 40:
                        content_started = True
                    continue
                if line.strip():
                    pdf.multi_cell(0, 7, line.strip())
                    pdf.ln(2)
        else:
            for line in body_text.split('\n'):
                if line.strip():
                    pdf.multi_cell(0, 7, line.strip())
                    pdf.ln(2)

        pdf.output(filepath)
        return filepath
